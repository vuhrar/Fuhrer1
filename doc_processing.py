# doc_processing.py
"""
استخراج النص من ملفات PDF/DOCX/TXT/HTML/CSV و
دوال استخراج مواد قانونية من النصوص.
"""
from typing import List, Dict
import io, re
from datetime import datetime
from .utils import _bytes, _norm


def extract_laws_from_text(text: str, source: str = "") -> List[Dict]:
    records = []
    # تقسيم يعتمد على بداية المادة العربية "المادة" أو كلمة "المادة" متبوعة
    parts = re.split(r'(?=المادة\s+)', text)
    current_law = "الأنظمة السعودية"
    for part in parts:
        part = part.strip()
        if len(part) < 40:
            continue
        lm = re.search(r'نظام\s+[\u0600-\u06ff\s]{4,35}(?=\n|،|\.)', part)
        if lm:
            current_law = lm.group(0).strip()
        am = re.match(r'(المادة\s+[\u0600-\u06ff\s\d]{2,40}?)(?::|[\n])', part)
        article = am.group(1).strip() if am else ""
        clean = re.sub(r'https?://\S+', '', part)
        clean = re.sub(r'\s+', ' ', clean).strip()
        if sum(1 for c in clean if '\u0600' <= c <= '\u06ff') < 20:
            continue
        records.append({
            "text": clean[:800],
            "article": article,
            "law_name": current_law,
            "source": source,
            "ts": datetime.now().strftime("%Y-%m-%d"),
        })
    return records


def extract_laws_from_pdf(raw: bytes, source: str = "") -> List[Dict]:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text() or ""
                if t.strip():
                    text += t + "\n"
    except Exception:
        try:
            import PyPDF2
            for pg in PyPDF2.PdfReader(io.BytesIO(raw)).pages:
                t = pg.extract_text() or ""
                if t.strip():
                    text += t + "\n"
        except Exception:
            pass
    return extract_laws_from_text(text, source) if text else []


def extract_laws_from_docx(raw: bytes, source: str = "") -> List[Dict]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return extract_laws_from_text(text, source)
    except Exception:
        return []


class DocIntel:
    def extract(self, f) -> str:
        ext = (getattr(f, "name", "") or "").rsplit(".", 1)[-1].lower()
        raw = _bytes(f)
        try:
            if ext == "pdf":
                return self._pdf(raw)
            if ext == "docx":
                return self._docx(raw)
            if ext in ("html", "htm"):
                return _norm(re.sub(r'<[^>]+>', '', raw.decode("utf-8", errors="ignore")))
            if ext == "json":
                return _norm(json.dumps(json.loads(raw.decode("utf-8", errors="ignore")), ensure_ascii=False))
            if ext == "csv":
                import csv
                rows = list(csv.reader(io.StringIO(raw.decode("utf-8", errors="ignore"))))
                return _norm("\n".join(" | ".join(r) for r in rows))
            return _norm(raw.decode("utf-8", errors="ignore"))
        except Exception:
            return ""

    def _pdf(self, raw: bytes) -> str:
        parts = []
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for pg in pdf.pages:
                    t = pg.extract_text() or ""
                    if t.strip():
                        parts.append(t)
        except Exception:
            try:
                import PyPDF2
                for pg in PyPDF2.PdfReader(io.BytesIO(raw)).pages:
                    t = pg.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            except Exception:
                pass
        return _norm("\n".join(parts))

    def _docx(self, raw: bytes) -> str:
        try:
            from docx import Document
            return _norm("\n".join(p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text))
        except Exception:
            return ""

    def entities(self, t: str) -> dict:
        return {
            "parties":  list(set(re.findall(r"(?:المدعي|المدعى عليه|الشركة|المؤسسة|الموظف|الهيئة)", t or ""))),
            "amounts":  re.findall(r"[\d,]+\s*(?:ريال|درهم|دولار)", t or ""),
            "articles": re.findall(r"المادة\s*[\u0600-\u06FF\d]+", t or ""),
            "dates":    re.findall(r"\d{1,2}/\d{1,2}/\d{2,4}", t or ""),
        }
